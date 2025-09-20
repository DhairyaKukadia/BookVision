from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fpdf import FPDF
import os
import logging

logger = logging.getLogger(__name__)

PDF_FONT_FAMILY: str = "Arial"
PDF_HEADING_SIZE: int = 16
PDF_SUBHEADING_SIZE: int = 14
PDF_TEXT_SIZE: int = 12
PDF_LINE_HEIGHT: int = 10

def save_as_docx(filename: str, extracted_text: str, summary: str, sentiment: str) -> bool:
    logger.info(f"Attempting to save formatted content to DOCX: {filename}")
    try:
        doc = Document()
        
        doc.add_heading('BookVision Analysis Report', level=0)

        doc.add_heading('Full Extracted Text', level=1)
        for paragraph_text in extracted_text.split('\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text)
            else:
                p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

        doc.add_heading('Summary', level=1)
        for paragraph_text in summary.split('\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text)
            else:
                p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()

        doc.add_heading('Sentiment Analysis', level=1)
        p = doc.add_paragraph()
        run = p.add_run(sentiment)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(6)

        doc.save(filename)
        logger.info(f"Formatted content successfully saved to DOCX: {filename}")
        return True
    except Exception as e:
        logger.error(f"Error saving formatted content to DOCX file {filename}: {e}", exc_info=True)
        return False

def save_as_pdf(filename: str, extracted_text: str, summary: str, sentiment: str) -> bool:
    logger.info(f"Attempting to save formatted content to PDF: {filename}")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font(PDF_FONT_FAMILY, 'B', PDF_HEADING_SIZE + 2)
        pdf.cell(0, PDF_LINE_HEIGHT * 1.5, 'BookVision Analysis Report', 0, 1, 'C')
        pdf.ln(PDF_LINE_HEIGHT)

        pdf.set_font(PDF_FONT_FAMILY, 'B', PDF_HEADING_SIZE)
        pdf.cell(0, PDF_LINE_HEIGHT, 'Full Extracted Text', 0, 1, 'L')
        pdf.set_font(PDF_FONT_FAMILY, '', PDF_TEXT_SIZE)
        try:
            encoded_extracted_text = extracted_text.encode('latin-1', 'replace').decode('latin-1')
        except Exception:
            encoded_extracted_text = extracted_text
        pdf.multi_cell(0, PDF_LINE_HEIGHT, txt=encoded_extracted_text)
        pdf.ln(PDF_LINE_HEIGHT * 0.5)

        pdf.set_font(PDF_FONT_FAMILY, 'B', PDF_HEADING_SIZE)
        pdf.cell(0, PDF_LINE_HEIGHT, 'Summary', 0, 1, 'L')
        pdf.set_font(PDF_FONT_FAMILY, '', PDF_TEXT_SIZE)
        try:
            encoded_summary = summary.encode('latin-1', 'replace').decode('latin-1')
        except Exception:
            encoded_summary = summary
        pdf.multi_cell(0, PDF_LINE_HEIGHT, txt=encoded_summary)
        pdf.ln(PDF_LINE_HEIGHT * 0.5)

        pdf.set_font(PDF_FONT_FAMILY, 'B', PDF_HEADING_SIZE)
        pdf.cell(0, PDF_LINE_HEIGHT, 'Sentiment Analysis', 0, 1, 'L')
        pdf.set_font(PDF_FONT_FAMILY, 'B', PDF_TEXT_SIZE)
        try:
            encoded_sentiment = sentiment.encode('latin-1', 'replace').decode('latin-1')
        except Exception:
            encoded_sentiment = sentiment
        pdf.multi_cell(0, PDF_LINE_HEIGHT, txt=encoded_sentiment)
        
        pdf.output(filename, "F")
        logger.info(f"Formatted content successfully saved to PDF: {filename}")
        return True
    except Exception as e:
        logger.error(f"Error saving formatted content to PDF file {filename}: {e}", exc_info=True)
        return False
